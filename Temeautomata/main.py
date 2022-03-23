import wikipedia
from docx import Document

language = input(str("Alege limba temei(ro,en,es...): "))
wikipedia.set_lang(language)
document = Document()

x = input("Despre ce subiect este tema? ")

titlu = wikipedia.page(x).title
continut = wikipedia.page(x).content
document.add_heading(titlu)
paragraph = document.add_paragraph(continut)
content = titlu + "\n" + continut

document.save('tema.docx')
